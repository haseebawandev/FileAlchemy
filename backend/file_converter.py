"""
Universal File Conversion Service
Supports images, documents, media, and archives
"""

import os
import subprocess
from pathlib import Path
from typing import Optional, List, Dict, Any
from abc import ABC, abstractmethod

class BaseConverter(ABC):
    """Base class for all file converters"""
    
    @abstractmethod
    def convert(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert file from input_path to output_path"""
        pass
    
    @abstractmethod
    def supported_formats(self) -> Dict[str, List[str]]:
        """Return dict of supported input->output format mappings"""
        pass

class ImageConverter(BaseConverter):
    """Handle image conversions using Pillow and pillow-heif"""
    
    def __init__(self):
        self.available_libs = {}
        
        # Check for basic image support (Pillow)
        try:
            from PIL import Image
            self.available_libs['pillow'] = True
        except ImportError as e:
            print(f"Basic image conversion unavailable: {e}")
            self.available_libs['pillow'] = False
        
        # Check for HEIF support (optional)
        try:
            import pillow_heif
            pillow_heif.register_heif_opener()
            self.available_libs['pillow_heif'] = True
        except ImportError as e:
            print(f"HEIF image support unavailable (this is optional): {e}")
            self.available_libs['pillow_heif'] = False
            
        # Check for SVG support (optional, requires system Cairo library)
        try:
            import cairosvg
            self.available_libs['cairosvg'] = True
        except (ImportError, OSError) as e:
            print(f"SVG conversion unavailable (this is optional): {e}")
            self.available_libs['cairosvg'] = False
            
        self.available = self.available_libs['pillow']
    
    def convert(self, input_path: str, output_path: str, **kwargs) -> bool:
        if not self.available:
            return False
            
        input_ext = Path(input_path).suffix.lower().lstrip('.')
        output_ext = Path(output_path).suffix.lower().lstrip('.')
        
        # Handle SVG conversion
        if input_ext == 'svg':
            return self._convert_svg(input_path, output_path, **kwargs)
            
        try:
            from PIL import Image
            
            with Image.open(input_path) as img:
                # Handle different color modes for JPEG conversion
                if output_ext in ['jpg', 'jpeg']:
                    if img.mode in ['RGBA', 'LA']:
                        # Create white background for JPEG (transparency not supported)
                        background = Image.new('RGB', img.size, (255, 255, 255))
                        background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                        img = background
                    elif img.mode == 'P':
                        # Convert palette mode (common in GIFs) to RGB
                        img = img.convert('RGB')
                    elif img.mode not in ['RGB', 'L']:
                        # Convert any other mode to RGB
                        img = img.convert('RGB')
                
                # Convert and save
                save_format = output_ext.upper()
                if save_format in ['JPG']:
                    save_format = 'JPEG'
                img.save(output_path, format=save_format, **kwargs)
            return True
            
        except Exception as e:
            print(f"Image conversion failed: {e}")
            return False
    
    def _convert_svg(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert SVG to raster format"""
        if not self.available_libs.get('cairosvg', False):
            print("cairosvg not available for SVG conversion")
            return False
            
        try:
            import cairosvg
            from PIL import Image
            import io
            
            output_ext = Path(output_path).suffix.lower().lstrip('.')
            
            if output_ext in ['png']:
                cairosvg.svg2png(url=input_path, write_to=output_path)
            elif output_ext in ['jpg', 'jpeg']:
                # Convert SVG to PNG first, then to JPEG
                png_data = cairosvg.svg2png(url=input_path)
                img = Image.open(io.BytesIO(png_data))
                # Convert RGBA to RGB for JPEG
                if img.mode == 'RGBA':
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[-1])
                    img = background
                img.save(output_path, 'JPEG')
            else:
                return False
                
            return True
        except Exception as e:
            print(f"SVG conversion failed: {e}")
            return False
    
    def supported_formats(self) -> Dict[str, List[str]]:
        formats = {
            'input': ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'gif', 'heic', 'heif', 'webp', 'ico'],
            'output': ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'gif', 'webp', 'ico']
        }
        # Add SVG support if cairosvg is available
        if self.available_libs.get('cairosvg', False):
            formats['input'].append('svg')
        return formats

class DocumentConverter(BaseConverter):
    """Handle document conversions using PyMuPDF and pdf2docx"""
    
    def __init__(self):
        self.available_libs = {}
        try:
            import fitz  # PyMuPDF
            self.available_libs['pymupdf'] = True
        except ImportError:
            self.available_libs['pymupdf'] = False
            
        try:
            from pdf2docx import Converter
            self.available_libs['pdf2docx'] = True
        except ImportError:
            self.available_libs['pdf2docx'] = False
            
        # Check for python-docx (for reading DOCX files)
        try:
            import docx
            self.available_libs['python_docx'] = True
        except ImportError:
            self.available_libs['python_docx'] = False
            
        # Check for reportlab (for creating PDFs)
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            self.available_libs['reportlab'] = True
        except ImportError:
            self.available_libs['reportlab'] = False
    
    def convert(self, input_path: str, output_path: str, **kwargs) -> bool:
        input_ext = Path(input_path).suffix.lower()
        output_ext = Path(output_path).suffix.lower()
        
        # PDF to DOCX conversion
        if input_ext == '.pdf' and output_ext == '.docx':
            return self._pdf_to_docx(input_path, output_path, **kwargs)
        
        # PDF to text conversion
        elif input_ext == '.pdf' and output_ext == '.txt':
            return self._pdf_to_text(input_path, output_path, **kwargs)
        
        # PDF to image conversion (multi-page to ZIP)
        elif input_ext == '.pdf' and output_ext in ['.jpg', '.jpeg', '.png']:
            return self._pdf_to_images(input_path, output_path, **kwargs)
        
        # DOCX to PDF conversion
        elif input_ext == '.docx' and output_ext == '.pdf':
            return self._docx_to_pdf(input_path, output_path, **kwargs)
        
        # DOCX to TXT conversion
        elif input_ext == '.docx' and output_ext == '.txt':
            return self._docx_to_txt(input_path, output_path, **kwargs)
        
        # TXT to DOCX conversion
        elif input_ext == '.txt' and output_ext == '.docx':
            return self._txt_to_docx(input_path, output_path, **kwargs)
        
        # TXT to PDF conversion
        elif input_ext == '.txt' and output_ext == '.pdf':
            return self._txt_to_pdf(input_path, output_path, **kwargs)
        
        # TXT to HTML conversion
        elif input_ext == '.txt' and output_ext == '.html':
            return self._txt_to_html(input_path, output_path, **kwargs)
        
        # HTML to TXT conversion
        elif input_ext == '.html' and output_ext == '.txt':
            return self._html_to_txt(input_path, output_path, **kwargs)
        
        # HTML to PDF conversion
        elif input_ext == '.html' and output_ext == '.pdf':
            return self._html_to_pdf(input_path, output_path, **kwargs)
        
        return False
    
    def _pdf_to_docx(self, input_path: str, output_path: str, **kwargs) -> bool:
        if not self.available_libs['pdf2docx']:
            print("pdf2docx not available")
            return False
            
        try:
            from pdf2docx import Converter
            cv = Converter(input_path)
            cv.convert(output_path, start=0, end=None)
            cv.close()
            return True
        except Exception as e:
            print(f"PDF to DOCX conversion failed: {e}")
            return False
    
    def _pdf_to_text(self, input_path: str, output_path: str, **kwargs) -> bool:
        if not self.available_libs['pymupdf']:
            print("PyMuPDF not available")
            return False
            
        try:
            import fitz
            doc = fitz.open(input_path)
            text = ""
            for page in doc:
                text += page.get_text()
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            doc.close()
            return True
        except Exception as e:
            print(f"PDF to text conversion failed: {e}")
            return False
    
    def _pdf_to_images(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert PDF pages to images and package in ZIP file"""
        print(f"Starting PDF to images conversion: {input_path} -> {output_path}")
        
        if not self.available_libs['pymupdf']:
            print("PyMuPDF not available for PDF to image conversion")
            return False
            
        try:
            import fitz
            import zipfile
            import tempfile
            import os
            
            # For PDF to images, we need to determine the target format from kwargs or default to jpg
            target_format = kwargs.get('target_format', 'jpg').lower()
            if target_format in ['jpeg']:
                target_format = 'jpg'
            
            print(f"Target image format: {target_format}")
            
            if target_format not in ['jpg', 'png']:
                print(f"Unsupported image format: {target_format}")
                return False
            
            # Open PDF document
            doc = fitz.open(input_path)
            
            # Create temporary directory for images
            with tempfile.TemporaryDirectory() as temp_dir:
                image_files = []
                
                # Convert each page to image
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    
                    # Set resolution (DPI) - higher values = better quality but larger files
                    dpi = kwargs.get('dpi', 150)  # Default 150 DPI
                    zoom = dpi / 72  # 72 is the default DPI
                    mat = fitz.Matrix(zoom, zoom)
                    
                    # Render page to pixmap
                    pix = page.get_pixmap(matrix=mat)
                    
                    # Generate filename for this page
                    base_name = Path(input_path).stem
                    page_filename = f"{base_name}_page_{page_num + 1:03d}.{target_format}"
                    page_path = os.path.join(temp_dir, page_filename)
                    
                    # Save image
                    if target_format in ['jpg', 'jpeg']:
                        # Convert to RGB for JPEG (remove alpha channel)
                        if pix.alpha:
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                        pix.save(page_path)
                    else:
                        # PNG supports alpha channel
                        pix.save(page_path)
                    
                    image_files.append((page_filename, page_path))
                    pix = None  # Free memory
                
                doc.close()
                
                # Create ZIP file with all images
                with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    for filename, filepath in image_files:
                        zipf.write(filepath, filename)
                
                print(f"Successfully converted {len(image_files)} pages to {target_format.upper()} images in ZIP file")
                return True
                
        except Exception as e:
            print(f"PDF to images conversion failed: {e}")
            return False
    
    def _docx_to_pdf(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert DOCX to PDF using python-docx and reportlab"""
        print(f"Starting DOCX to PDF conversion: {input_path} -> {output_path}")
        
        if not self.available_libs['python_docx']:
            print("python-docx not available for DOCX reading")
            return False
            
        if not self.available_libs['reportlab']:
            print("reportlab not available for PDF creation")
            return False
            
        try:
            import docx
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.units import inch
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib import colors
            import io
            
            # Read DOCX document
            doc = docx.Document(input_path)
            
            # Create PDF document
            pdf_doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            story = []
            
            # Process each paragraph in the DOCX
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Skip empty paragraphs
                    # Determine style based on paragraph formatting
                    if paragraph.style.name.startswith('Heading'):
                        # Use heading style
                        if '1' in paragraph.style.name:
                            style = styles['Heading1']
                        elif '2' in paragraph.style.name:
                            style = styles['Heading2']
                        else:
                            style = styles['Heading3']
                    else:
                        # Use normal style
                        style = styles['Normal']
                    
                    # Create paragraph with text
                    para = Paragraph(paragraph.text, style)
                    story.append(para)
                    story.append(Spacer(1, 12))  # Add space between paragraphs
            
            # Handle tables if present
            for table in doc.tables:
                # Add table content as text (simple approach)
                story.append(Paragraph("--- Table Content ---", styles['Heading3']))
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        para = Paragraph(row_text, styles['Normal'])
                        story.append(para)
                        story.append(Spacer(1, 6))
                story.append(Spacer(1, 12))
            
            # Build PDF
            pdf_doc.build(story)
            
            print(f"Successfully converted DOCX to PDF: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
            return True
            
        except Exception as e:
            print(f"DOCX to PDF conversion failed: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _docx_to_txt(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert DOCX to plain text"""
        print(f"Starting DOCX to TXT conversion: {input_path} -> {output_path}")
        
        if not self.available_libs['python_docx']:
            print("python-docx not available for DOCX reading")
            return False
            
        try:
            import docx
            
            # Read DOCX document
            doc = docx.Document(input_path)
            
            # Extract all text content
            text_content = []
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Process tables
            for table in doc.tables:
                text_content.append("\n--- Table Content ---")
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text_content.append(row_text)
                text_content.append("--- End Table ---\n")
            
            # Write to text file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(text_content))
            
            print(f"Successfully converted DOCX to TXT: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
            return True
            
        except Exception as e:
            print(f"DOCX to TXT conversion failed: {e}")
            return False
    
    def _txt_to_docx(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert plain text to DOCX"""
        print(f"Starting TXT to DOCX conversion: {input_path} -> {output_path}")
        
        if not self.available_libs['python_docx']:
            print("python-docx not available for DOCX creation")
            return False
            
        try:
            import docx
            
            # Read text file
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Create new DOCX document
            doc = docx.Document()
            
            # Add title based on filename
            filename = Path(input_path).stem
            doc.add_heading(f'Document: {filename}', 0)
            
            # Split content into paragraphs
            paragraphs = content.split('\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Check if it looks like a heading (all caps, short, etc.)
                    if (len(para_text.strip()) < 50 and 
                        para_text.strip().isupper() and 
                        not para_text.strip().endswith('.')):
                        doc.add_heading(para_text.strip(), level=1)
                    else:
                        doc.add_paragraph(para_text)
                else:
                    # Add empty paragraph for spacing
                    doc.add_paragraph('')
            
            # Save DOCX document
            doc.save(output_path)
            
            print(f"Successfully converted TXT to DOCX: {len(paragraphs)} paragraphs")
            return True
            
        except Exception as e:
            print(f"TXT to DOCX conversion failed: {e}")
            return False
    
    def _txt_to_pdf(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert plain text to PDF"""
        print(f"Starting TXT to PDF conversion: {input_path} -> {output_path}")
        
        if not self.available_libs['reportlab']:
            print("reportlab not available for PDF creation")
            return False
            
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import inch
            
            # Read text file
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Create PDF document
            pdf_doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            story = []
            
            # Add title
            filename = Path(input_path).stem
            title = Paragraph(f"Text Document: {filename}", styles['Title'])
            story.append(title)
            story.append(Spacer(1, 20))
            
            # Split content into paragraphs
            paragraphs = content.split('\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Check if it looks like a heading
                    if (len(para_text.strip()) < 50 and 
                        para_text.strip().isupper() and 
                        not para_text.strip().endswith('.')):
                        para = Paragraph(para_text.strip(), styles['Heading1'])
                    else:
                        para = Paragraph(para_text, styles['Normal'])
                    story.append(para)
                    story.append(Spacer(1, 12))
                else:
                    story.append(Spacer(1, 12))
            
            # Build PDF
            pdf_doc.build(story)
            
            print(f"Successfully converted TXT to PDF: {len(paragraphs)} paragraphs")
            return True
            
        except Exception as e:
            print(f"TXT to PDF conversion failed: {e}")
            return False
    
    def _txt_to_html(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert plain text to HTML"""
        print(f"Starting TXT to HTML conversion: {input_path} -> {output_path}")
        
        try:
            # Read text file
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Get filename for title
            filename = Path(input_path).stem
            
            # Create HTML content
            html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{filename}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f9f9f9;
        }}
        .container {{
            background-color: white;
            padding: 30px;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #333;
            border-bottom: 2px solid #007acc;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #555;
            margin-top: 30px;
        }}
        p {{
            margin-bottom: 15px;
            text-align: justify;
        }}
        .empty-line {{
            height: 15px;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>Text Document: {filename}</h1>
"""
            
            # Process content
            paragraphs = content.split('\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Check if it looks like a heading
                    if (len(para_text.strip()) < 50 and 
                        para_text.strip().isupper() and 
                        not para_text.strip().endswith('.')):
                        html_content += f"        <h2>{para_text.strip()}</h2>\n"
                    else:
                        # Escape HTML characters and preserve line breaks
                        escaped_text = para_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        html_content += f"        <p>{escaped_text}</p>\n"
                else:
                    html_content += '        <div class="empty-line"></div>\n'
            
            # Close HTML
            html_content += """    </div>
</body>
</html>"""
            
            # Write HTML file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            print(f"Successfully converted TXT to HTML: {len(paragraphs)} paragraphs")
            return True
            
        except Exception as e:
            print(f"TXT to HTML conversion failed: {e}")
            return False
    
    def _html_to_txt(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert HTML to plain text"""
        print(f"Starting HTML to TXT conversion: {input_path} -> {output_path}")
        
        try:
            # Try to use BeautifulSoup if available, otherwise use simple parsing
            try:
                from bs4 import BeautifulSoup
                
                # Read HTML file
                with open(input_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                
                # Parse HTML
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Extract text content
                text_content = soup.get_text()
                
                # Clean up extra whitespace
                lines = [line.strip() for line in text_content.split('\n')]
                cleaned_lines = []
                
                for line in lines:
                    if line:
                        cleaned_lines.append(line)
                    elif cleaned_lines and cleaned_lines[-1]:  # Add empty line only if previous line had content
                        cleaned_lines.append('')
                
                text_content = '\n'.join(cleaned_lines)
                
            except ImportError:
                print("BeautifulSoup not available, using simple HTML parsing")
                
                # Simple HTML parsing without BeautifulSoup
                with open(input_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                
                # Remove HTML tags using simple regex
                import re
                
                # Remove script and style elements
                html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
                html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
                
                # Remove HTML tags
                text_content = re.sub(r'<[^>]+>', '', html_content)
                
                # Decode HTML entities
                text_content = text_content.replace('&amp;', '&')
                text_content = text_content.replace('&lt;', '<')
                text_content = text_content.replace('&gt;', '>')
                text_content = text_content.replace('&quot;', '"')
                text_content = text_content.replace('&#39;', "'")
                text_content = text_content.replace('&nbsp;', ' ')
                
                # Clean up whitespace
                lines = [line.strip() for line in text_content.split('\n')]
                text_content = '\n'.join(line for line in lines if line)
            
            # Write text file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            print(f"Successfully converted HTML to TXT")
            return True
            
        except Exception as e:
            print(f"HTML to TXT conversion failed: {e}")
            return False
    
    def _html_to_pdf(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert HTML to PDF"""
        print(f"Starting HTML to PDF conversion: {input_path} -> {output_path}")
        
        if not self.available_libs['reportlab']:
            print("reportlab not available for PDF creation")
            return False
            
        try:
            # First convert HTML to text, then text to PDF
            import tempfile
            
            with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as temp_file:
                temp_txt_path = temp_file.name
            
            # Convert HTML to text first
            if not self._html_to_txt(input_path, temp_txt_path):
                return False
            
            # Then convert text to PDF
            result = self._txt_to_pdf(temp_txt_path, output_path, **kwargs)
            
            # Cleanup temp file
            import os
            try:
                os.unlink(temp_txt_path)
            except:
                pass
            
            if result:
                print(f"Successfully converted HTML to PDF via text conversion")
            
            return result
            
        except Exception as e:
            print(f"HTML to PDF conversion failed: {e}")
            return False
    
    def supported_formats(self) -> Dict[str, List[str]]:
        formats = {'input': [], 'output': []}
        
        # PDF conversions
        if self.available_libs['pymupdf']:
            formats['input'].extend(['pdf'])
            formats['output'].extend(['txt', 'jpg', 'jpeg', 'png'])
        if self.available_libs['pdf2docx']:
            formats['output'].extend(['docx'])
        
        # DOCX conversions
        if self.available_libs['python_docx']:
            formats['input'].extend(['docx'])
            formats['output'].extend(['txt'])
            if self.available_libs['reportlab']:
                formats['output'].extend(['pdf'])
        
        # Text conversions
        formats['input'].extend(['txt', 'html'])  # Always available
        formats['output'].extend(['html'])  # Always available
        
        if self.available_libs['python_docx']:
            formats['output'].extend(['docx'])
        if self.available_libs['reportlab']:
            formats['output'].extend(['pdf'])
        
        # Remove duplicates
        formats['input'] = list(set(formats['input']))
        formats['output'] = list(set(formats['output']))
        
        return formats

class MediaConverter(BaseConverter):
    """Handle video/audio conversions using FFmpeg"""
    
    def __init__(self):
        self.available = self._check_ffmpeg()
    
    def _check_ffmpeg(self) -> bool:
        try:
            subprocess.run(['ffmpeg', '-version'], 
                         capture_output=True, check=True)
            return True
        except (subprocess.CalledProcessError, FileNotFoundError):
            print("FFmpeg not found. Install FFmpeg for media conversion.")
            return False
    
    def convert(self, input_path: str, output_path: str, **kwargs) -> bool:
        if not self.available:
            return False
            
        input_ext = Path(input_path).suffix.lower().lstrip('.')
        output_ext = Path(output_path).suffix.lower().lstrip('.')
        
        print(f"Starting media conversion: {input_ext} -> {output_ext}")
        
        try:
            cmd = ['ffmpeg', '-i', input_path]
            
            # Special handling for different conversion types
            if input_ext == 'gif' and output_ext in ['mp4', 'avi', 'mov', 'mkv', 'webm']:
                # GIF to video conversion
                cmd.extend(['-movflags', 'faststart'])
                cmd.extend(['-pix_fmt', 'yuv420p'])
                if 'fps' in kwargs:
                    cmd.extend(['-r', str(kwargs['fps'])])
                else:
                    cmd.extend(['-r', '15'])  # Default 15 fps for GIF to video
                    
            elif output_ext == 'gif':
                # Video to GIF conversion
                # Use palette generation for better quality
                if input_ext in ['mp4', 'avi', 'mov', 'mkv', 'webm']:
                    # Two-pass conversion for better GIF quality
                    palette_path = output_path.replace('.gif', '_palette.png')
                    
                    # First pass: generate palette
                    palette_cmd = ['ffmpeg', '-i', input_path, '-vf', 'palettegen', '-y', palette_path]
                    palette_result = subprocess.run(palette_cmd, capture_output=True, text=True)
                    
                    if palette_result.returncode == 0:
                        # Second pass: create GIF with palette
                        fps = kwargs.get('fps', 10)  # Default 10 fps for video to GIF
                        scale = kwargs.get('scale', '320:-1')  # Default scale
                        cmd = ['ffmpeg', '-i', input_path, '-i', palette_path]
                        cmd.extend(['-filter_complex', f'fps={fps},scale={scale}:flags=lanczos[x];[x][1:v]paletteuse'])
                        cmd.extend(['-y', output_path])
                        
                        result = subprocess.run(cmd, capture_output=True, text=True)
                        
                        # Cleanup palette file
                        try:
                            os.remove(palette_path)
                        except:
                            pass
                        
                        if result.returncode == 0:
                            print(f"Successfully converted {input_ext} to GIF with palette optimization")
                            return True
                    
                    # Fallback to simple conversion if palette method fails
                    print("Palette method failed, using simple GIF conversion")
                    cmd = ['ffmpeg', '-i', input_path]
                    fps = kwargs.get('fps', 10)
                    scale = kwargs.get('scale', '320:-1')
                    cmd.extend(['-vf', f'fps={fps},scale={scale}:flags=lanczos'])
                    
            else:
                # Standard video/audio conversion
                if 'quality' in kwargs:
                    cmd.extend(['-crf', str(kwargs['quality'])])
                if 'bitrate' in kwargs:
                    cmd.extend(['-b:v', kwargs['bitrate']])
                if 'audio_bitrate' in kwargs:
                    cmd.extend(['-b:a', kwargs['audio_bitrate']])
            
            cmd.extend(['-y', output_path])  # -y to overwrite
            
            print(f"Running FFmpeg command: {' '.join(cmd[:5])}... (truncated)")
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode == 0:
                print(f"Media conversion successful: {input_ext} -> {output_ext}")
                return True
            else:
                print(f"FFmpeg error: {result.stderr}")
                return False
            
        except Exception as e:
            print(f"Media conversion failed: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def supported_formats(self) -> Dict[str, List[str]]:
        return {
            'input': ['mp4', 'avi', 'mov', 'mkv', 'wmv', 'flv', 'webm', 'gif', 'mp3', 'wav', 'flac', 'aac', 'ogg'],
            'output': ['mp4', 'avi', 'mov', 'mkv', 'webm', 'wmv', 'gif', 'mp3', 'wav', 'aac', 'flac', 'ogg']
        }

class ArchiveConverter(BaseConverter):
    """Handle archive conversions using various archive libraries"""
    
    def __init__(self):
        self.available_libs = {}
        
        # Check for built-in zipfile
        try:
            import zipfile
            self.available_libs['zipfile'] = True
        except ImportError:
            self.available_libs['zipfile'] = False
            
        # Check for tarfile (built-in)
        try:
            import tarfile
            self.available_libs['tarfile'] = True
        except ImportError:
            self.available_libs['tarfile'] = False
            
        # Check for py7zr
        try:
            import py7zr
            self.available_libs['py7zr'] = True
        except ImportError:
            self.available_libs['py7zr'] = False
            
        # Check for rarfile
        try:
            import rarfile
            self.available_libs['rarfile'] = True
        except ImportError:
            self.available_libs['rarfile'] = False
    
    def convert(self, input_path: str, output_path: str, **kwargs) -> bool:
        input_ext = Path(input_path).suffix.lower().lstrip('.')
        output_ext = Path(output_path).suffix.lower().lstrip('.')
        
        try:
            # Extract input archive
            temp_dir = Path("temp_extract")
            temp_dir.mkdir(exist_ok=True)
            
            if not self._extract_archive(input_path, input_ext, temp_dir):
                return False
            
            # Create output archive
            success = self._create_archive(temp_dir, output_path, output_ext)
            
            # Cleanup
            import shutil
            shutil.rmtree(temp_dir, ignore_errors=True)
            
            return success
            
        except Exception as e:
            print(f"Archive conversion failed: {e}")
            return False
    
    def _extract_archive(self, archive_path: str, format_type: str, extract_to: Path) -> bool:
        """Extract archive to temporary directory"""
        try:
            if format_type == 'zip' and self.available_libs['zipfile']:
                import zipfile
                with zipfile.ZipFile(archive_path, 'r') as zip_ref:
                    zip_ref.extractall(extract_to)
                return True
                
            elif format_type in ['tar', 'gz', 'tgz'] and self.available_libs['tarfile']:
                import tarfile
                with tarfile.open(archive_path, 'r:*') as tar_ref:
                    tar_ref.extractall(extract_to)
                return True
                
            elif format_type == '7z' and self.available_libs['py7zr']:
                import py7zr
                with py7zr.SevenZipFile(archive_path, mode='r') as z:
                    z.extractall(extract_to)
                return True
                
            elif format_type == 'rar' and self.available_libs['rarfile']:
                import rarfile
                with rarfile.RarFile(archive_path) as rf:
                    rf.extractall(extract_to)
                return True
                
            return False
        except Exception as e:
            print(f"Failed to extract {format_type}: {e}")
            return False
    
    def _create_archive(self, source_dir: Path, output_path: str, format_type: str) -> bool:
        """Create archive from directory"""
        try:
            if format_type == 'zip' and self.available_libs['zipfile']:
                import zipfile
                with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    for file_path in source_dir.rglob('*'):
                        if file_path.is_file():
                            arcname = file_path.relative_to(source_dir)
                            zipf.write(file_path, arcname)
                return True
                
            elif format_type in ['tar', 'gz'] and self.available_libs['tarfile']:
                import tarfile
                mode = 'w:gz' if format_type == 'gz' else 'w'
                with tarfile.open(output_path, mode) as tar:
                    tar.add(source_dir, arcname='.')
                return True
                
            elif format_type == '7z' and self.available_libs['py7zr']:
                import py7zr
                with py7zr.SevenZipFile(output_path, 'w') as archive:
                    for file_path in source_dir.rglob('*'):
                        if file_path.is_file():
                            arcname = file_path.relative_to(source_dir)
                            archive.write(file_path, arcname)
                return True
                
            return False
        except Exception as e:
            print(f"Failed to create {format_type}: {e}")
            return False
    
    def supported_formats(self) -> Dict[str, List[str]]:
        formats = {'input': [], 'output': []}
        
        if self.available_libs['zipfile']:
            formats['input'].append('zip')
            formats['output'].append('zip')
            
        if self.available_libs['tarfile']:
            formats['input'].extend(['tar', 'gz', 'tgz'])
            formats['output'].extend(['tar', 'gz'])
            
        if self.available_libs['py7zr']:
            formats['input'].append('7z')
            formats['output'].append('7z')
            
        if self.available_libs['rarfile']:
            formats['input'].append('rar')
            # Note: rarfile can only extract, not create RAR files
            # RAR creation requires proprietary WinRAR software
            
        return formats

class DataConverter(BaseConverter):
    """Handle data file conversions (CSV, XLSX, JSON, etc.)"""
    
    def __init__(self):
        self.available_libs = {}
        
        # Check for pandas (for CSV/Excel operations)
        try:
            import pandas as pd
            self.available_libs['pandas'] = True
        except ImportError:
            self.available_libs['pandas'] = False
            
        # Check for openpyxl (for Excel files)
        try:
            import openpyxl
            self.available_libs['openpyxl'] = True
        except ImportError:
            self.available_libs['openpyxl'] = False
            
        # Check for reportlab (for PDF creation)
        try:
            from reportlab.pdfgen import canvas
            self.available_libs['reportlab'] = True
        except ImportError:
            self.available_libs['reportlab'] = False
    
    def convert(self, input_path: str, output_path: str, **kwargs) -> bool:
        input_ext = Path(input_path).suffix.lower()
        output_ext = Path(output_path).suffix.lower()
        
        # CSV conversions
        if input_ext == '.csv' and output_ext == '.xlsx':
            return self._csv_to_xlsx(input_path, output_path, **kwargs)
        elif input_ext == '.csv' and output_ext == '.pdf':
            return self._csv_to_pdf(input_path, output_path, **kwargs)
        elif input_ext == '.csv' and output_ext == '.txt':
            return self._csv_to_txt(input_path, output_path, **kwargs)
        elif input_ext == '.csv' and output_ext == '.json':
            return self._csv_to_json(input_path, output_path, **kwargs)
        
        # XLSX conversions
        elif input_ext == '.xlsx' and output_ext == '.csv':
            return self._xlsx_to_csv(input_path, output_path, **kwargs)
        elif input_ext == '.xlsx' and output_ext == '.pdf':
            return self._xlsx_to_pdf(input_path, output_path, **kwargs)
        elif input_ext == '.xlsx' and output_ext == '.txt':
            return self._xlsx_to_txt(input_path, output_path, **kwargs)
        
        # JSON conversions
        elif input_ext == '.json' and output_ext == '.csv':
            return self._json_to_csv(input_path, output_path, **kwargs)
        elif input_ext == '.json' and output_ext == '.xlsx':
            return self._json_to_xlsx(input_path, output_path, **kwargs)
        elif input_ext == '.json' and output_ext == '.txt':
            return self._json_to_txt(input_path, output_path, **kwargs)
        
        # XML conversions (basic support)
        elif input_ext == '.xml' and output_ext == '.json':
            return self._xml_to_json(input_path, output_path, **kwargs)
        elif input_ext == '.xml' and output_ext == '.txt':
            return self._xml_to_txt(input_path, output_path, **kwargs)
        
        return False
    
    def _csv_to_xlsx(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert CSV to Excel (XLSX)"""
        print(f"Starting CSV to XLSX conversion: {input_path} -> {output_path}")
        
        if not self.available_libs['pandas']:
            print("pandas not available for CSV processing")
            return False
            
        try:
            import pandas as pd
            
            # Read CSV file
            df = pd.read_csv(input_path, encoding='utf-8')
            
            # Write to Excel
            with pd.ExcelWriter(output_path, engine='openpyxl' if self.available_libs['openpyxl'] else 'xlsxwriter') as writer:
                df.to_excel(writer, sheet_name='Sheet1', index=False)
            
            print(f"Successfully converted CSV to XLSX: {len(df)} rows, {len(df.columns)} columns")
            return True
            
        except Exception as e:
            print(f"CSV to XLSX conversion failed: {e}")
            return False
    
    def _csv_to_pdf(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert CSV to PDF table"""
        print(f"Starting CSV to PDF conversion: {input_path} -> {output_path}")
        
        if not self.available_libs['pandas']:
            print("pandas not available for CSV processing")
            return False
            
        if not self.available_libs['reportlab']:
            print("reportlab not available for PDF creation")
            return False
            
        try:
            import pandas as pd
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            
            # Read CSV file
            df = pd.read_csv(input_path, encoding='utf-8')
            
            # Create PDF document
            pdf_doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=30,
                leftMargin=30,
                topMargin=30,
                bottomMargin=18
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            story = []
            
            # Add title
            filename = Path(input_path).stem
            title = Paragraph(f"Data Table: {filename}", styles['Title'])
            story.append(title)
            story.append(Spacer(1, 20))
            
            # Prepare table data
            # Convert DataFrame to list of lists for ReportLab
            table_data = []
            
            # Add headers
            headers = list(df.columns)
            table_data.append(headers)
            
            # Add data rows (limit to prevent huge PDFs)
            max_rows = kwargs.get('max_rows', 100)  # Default limit
            for i, row in df.iterrows():
                if i >= max_rows:
                    break
                # Convert all values to strings and handle NaN
                row_data = [str(val) if pd.notna(val) else '' for val in row.values]
                table_data.append(row_data)
            
            # Create table
            table = Table(table_data)
            
            # Style the table
            table.setStyle(TableStyle([
                # Header styling
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                
                # Data styling
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                
                # Alternating row colors
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.beige, colors.white]),
            ]))
            
            story.append(table)
            
            # Add summary info
            if len(df) > max_rows:
                story.append(Spacer(1, 20))
                summary = Paragraph(f"Note: Showing first {max_rows} rows of {len(df)} total rows", styles['Normal'])
                story.append(summary)
            
            # Build PDF
            pdf_doc.build(story)
            
            print(f"Successfully converted CSV to PDF: {len(df)} rows, {len(df.columns)} columns")
            return True
            
        except Exception as e:
            print(f"CSV to PDF conversion failed: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _csv_to_txt(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert CSV to formatted text"""
        print(f"Starting CSV to TXT conversion: {input_path} -> {output_path}")
        
        if not self.available_libs['pandas']:
            # Fallback to simple CSV reading
            try:
                import csv
                
                with open(input_path, 'r', encoding='utf-8') as csvfile:
                    reader = csv.reader(csvfile)
                    
                    with open(output_path, 'w', encoding='utf-8') as txtfile:
                        for row in reader:
                            txtfile.write('\t'.join(row) + '\n')
                
                print("Successfully converted CSV to TXT using basic CSV reader")
                return True
                
            except Exception as e:
                print(f"CSV to TXT conversion failed: {e}")
                return False
        
        try:
            import pandas as pd
            
            # Read CSV file
            df = pd.read_csv(input_path, encoding='utf-8')
            
            # Convert to formatted text
            with open(output_path, 'w', encoding='utf-8') as f:
                # Write header
                f.write("Data Table: " + Path(input_path).stem + "\n")
                f.write("=" * 50 + "\n\n")
                
                # Write data in a readable format
                f.write(df.to_string(index=False))
                
                # Write summary
                f.write(f"\n\nSummary: {len(df)} rows, {len(df.columns)} columns")
            
            print(f"Successfully converted CSV to TXT: {len(df)} rows, {len(df.columns)} columns")
            return True
            
        except Exception as e:
            print(f"CSV to TXT conversion failed: {e}")
            return False
    
    def _csv_to_json(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert CSV to JSON"""
        print(f"Starting CSV to JSON conversion: {input_path} -> {output_path}")
        
        if not self.available_libs['pandas']:
            # Fallback to basic CSV/JSON
            try:
                import csv
                import json
                
                data = []
                with open(input_path, 'r', encoding='utf-8') as csvfile:
                    reader = csv.DictReader(csvfile)
                    for row in reader:
                        data.append(row)
                
                with open(output_path, 'w', encoding='utf-8') as jsonfile:
                    json.dump(data, jsonfile, indent=2, ensure_ascii=False)
                
                print(f"Successfully converted CSV to JSON: {len(data)} records")
                return True
                
            except Exception as e:
                print(f"CSV to JSON conversion failed: {e}")
                return False
        
        try:
            import pandas as pd
            
            # Read CSV file
            df = pd.read_csv(input_path, encoding='utf-8')
            
            # Convert to JSON
            df.to_json(output_path, orient='records', indent=2, force_ascii=False)
            
            print(f"Successfully converted CSV to JSON: {len(df)} records")
            return True
            
        except Exception as e:
            print(f"CSV to JSON conversion failed: {e}")
            return False
    
    def _xlsx_to_csv(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert Excel (XLSX) to CSV"""
        print(f"Starting XLSX to CSV conversion: {input_path} -> {output_path}")
        
        if not self.available_libs['pandas']:
            print("pandas not available for XLSX processing")
            return False
            
        try:
            import pandas as pd
            
            # Read Excel file (first sheet by default)
            sheet_name = kwargs.get('sheet_name', 0)  # Default to first sheet
            df = pd.read_excel(input_path, sheet_name=sheet_name)
            
            # Write to CSV
            df.to_csv(output_path, index=False, encoding='utf-8')
            
            print(f"Successfully converted XLSX to CSV: {len(df)} rows, {len(df.columns)} columns")
            return True
            
        except Exception as e:
            print(f"XLSX to CSV conversion failed: {e}")
            return False
    
    def _xlsx_to_pdf(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert Excel (XLSX) to PDF"""
        print(f"Starting XLSX to PDF conversion: {input_path} -> {output_path}")
        
        # First convert to CSV, then CSV to PDF
        import tempfile
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False, encoding='utf-8') as temp_file:
            temp_csv_path = temp_file.name
        
        try:
            # Convert XLSX to CSV first
            if not self._xlsx_to_csv(input_path, temp_csv_path, **kwargs):
                return False
            
            # Then convert CSV to PDF
            result = self._csv_to_pdf(temp_csv_path, output_path, **kwargs)
            
            return result
            
        finally:
            # Cleanup temp file
            import os
            try:
                os.unlink(temp_csv_path)
            except:
                pass
    
    def _xlsx_to_txt(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert Excel (XLSX) to text"""
        print(f"Starting XLSX to TXT conversion: {input_path} -> {output_path}")
        
        # First convert to CSV, then CSV to TXT
        import tempfile
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False, encoding='utf-8') as temp_file:
            temp_csv_path = temp_file.name
        
        try:
            # Convert XLSX to CSV first
            if not self._xlsx_to_csv(input_path, temp_csv_path, **kwargs):
                return False
            
            # Then convert CSV to TXT
            result = self._csv_to_txt(temp_csv_path, output_path, **kwargs)
            
            return result
            
        finally:
            # Cleanup temp file
            import os
            try:
                os.unlink(temp_csv_path)
            except:
                pass
    
    def _json_to_csv(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert JSON to CSV"""
        print(f"Starting JSON to CSV conversion: {input_path} -> {output_path}")
        
        if not self.available_libs['pandas']:
            # Fallback to basic JSON/CSV
            try:
                import json
                import csv
                
                with open(input_path, 'r', encoding='utf-8') as jsonfile:
                    data = json.load(jsonfile)
                
                # Handle different JSON structures
                if isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict):
                    # List of dictionaries - standard tabular data
                    fieldnames = data[0].keys()
                    
                    with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
                        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                        writer.writeheader()
                        writer.writerows(data)
                    
                    print(f"Successfully converted JSON to CSV: {len(data)} records")
                    return True
                else:
                    print("JSON structure not suitable for CSV conversion")
                    return False
                    
            except Exception as e:
                print(f"JSON to CSV conversion failed: {e}")
                return False
        
        try:
            import pandas as pd
            
            # Read JSON file
            df = pd.read_json(input_path)
            
            # Write to CSV
            df.to_csv(output_path, index=False, encoding='utf-8')
            
            print(f"Successfully converted JSON to CSV: {len(df)} rows, {len(df.columns)} columns")
            return True
            
        except Exception as e:
            print(f"JSON to CSV conversion failed: {e}")
            return False
    
    def _json_to_xlsx(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert JSON to Excel (XLSX)"""
        print(f"Starting JSON to XLSX conversion: {input_path} -> {output_path}")
        
        # First convert to CSV, then CSV to XLSX
        import tempfile
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False, encoding='utf-8') as temp_file:
            temp_csv_path = temp_file.name
        
        try:
            # Convert JSON to CSV first
            if not self._json_to_csv(input_path, temp_csv_path, **kwargs):
                return False
            
            # Then convert CSV to XLSX
            result = self._csv_to_xlsx(temp_csv_path, output_path, **kwargs)
            
            return result
            
        finally:
            # Cleanup temp file
            import os
            try:
                os.unlink(temp_csv_path)
            except:
                pass
    
    def supported_formats(self) -> Dict[str, List[str]]:
        formats = {'input': [], 'output': []}
        
        if self.available_libs['pandas']:
            # CSV support
            formats['input'].extend(['csv'])
            formats['output'].extend(['txt', 'json'])
            
            if self.available_libs['openpyxl']:
                formats['input'].extend(['xlsx'])
                formats['output'].extend(['csv', 'xlsx'])
            
            if self.available_libs['reportlab']:
                formats['output'].extend(['pdf'])
        
        # JSON support (always available)
        formats['input'].extend(['json'])
        if self.available_libs['pandas']:
            formats['output'].extend(['csv'])
        
        # Remove duplicates
        formats['input'] = list(set(formats['input']))
        formats['output'] = list(set(formats['output']))
        
        return formats

class FileConversionService:
    """Main service class that orchestrates all converters"""
    
    def __init__(self):
        self.converters = {
            'image': ImageConverter(),
            'document': DocumentConverter(),
            'media': MediaConverter(),
            'archive': ArchiveConverter(),
            'data': DataConverter()
        }
    
    def convert_file(self, input_path: str, output_path: str, **kwargs) -> bool:
        """Convert a file based on its extension"""
        try:
            input_ext = Path(input_path).suffix.lower().lstrip('.')
            output_ext = Path(output_path).suffix.lower().lstrip('.')
            
            # Special case: PDF to image conversion where output file is .zip but target format is image
            target_format = kwargs.get('target_format')
            if input_ext == 'pdf' and output_ext == 'zip' and target_format:
                actual_target_ext = target_format.lower()
                print(f"PDF to image conversion detected: {input_ext} -> {actual_target_ext} (packaged as ZIP)")
                output_ext = actual_target_ext  # Use the actual target format for converter selection
            
            print(f"Converting: {input_ext} -> {output_ext}")
            print(f"Input file: {input_path}")
            print(f"Output file: {output_path}")
            
            # Check if input file exists
            if not os.path.exists(input_path):
                print(f"Error: Input file does not exist: {input_path}")
                return False
            
            # Check file size
            file_size = os.path.getsize(input_path)
            print(f"Input file size: {file_size} bytes")
            
            # Determine converter type
            converter_type = self._get_converter_type(input_ext, output_ext)
            if not converter_type:
                print(f"No converter found for {input_ext} -> {output_ext}")
                return False
            
            print(f"Using converter: {converter_type}")
            converter = self.converters[converter_type]
            
            # Perform conversion
            result = converter.convert(input_path, output_path, **kwargs)
            print(f"Conversion result: {result}")
            
            # Check if output file was created
            if result and os.path.exists(output_path):
                output_size = os.path.getsize(output_path)
                print(f"Output file created successfully, size: {output_size} bytes")
            elif result:
                print("Warning: Conversion reported success but output file not found")
                return False
            
            return result
            
        except Exception as e:
            print(f"Error in convert_file: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _get_converter_type(self, input_ext: str, output_ext: str) -> Optional[str]:
        """Determine which converter to use based on file extensions"""
        # Special case: PDF to image conversion (output will be ZIP, but we check for image format)
        if input_ext == 'pdf' and output_ext in ['jpg', 'jpeg', 'png']:
            return 'document'  # DocumentConverter handles PDF to images
        
        for conv_type, converter in self.converters.items():
            formats = converter.supported_formats()
            if input_ext in formats['input'] and output_ext in formats['output']:
                return conv_type
        return None
    
    def is_conversion_supported(self, input_ext: str, output_ext: str) -> tuple[bool, str]:
        """Check if a conversion is supported and return reason if not"""
        input_ext = input_ext.lower()
        output_ext = output_ext.lower()
        
        # Special case for RAR creation
        if output_ext == 'rar':
            return False, "RAR creation requires proprietary WinRAR software and is not supported"
        
        # Check if conversion is supported
        converter_type = self._get_converter_type(input_ext, output_ext)
        if converter_type:
            return True, "Conversion supported"
        
        # Check if input format is supported at all
        input_supported = False
        output_supported = False
        
        for conv_type, converter in self.converters.items():
            formats = converter.supported_formats()
            if input_ext in formats['input']:
                input_supported = True
            if output_ext in formats['output']:
                output_supported = True
        
        if not input_supported:
            return False, f"Input format '{input_ext}' is not supported"
        elif not output_supported:
            return False, f"Output format '{output_ext}' is not supported"
        else:
            return False, f"Conversion from '{input_ext}' to '{output_ext}' is not supported"
    
    def list_supported_formats(self) -> Dict[str, Dict[str, List[str]]]:
        """List all supported formats by converter type"""
        return {name: conv.supported_formats() 
                for name, conv in self.converters.items()}

if __name__ == "__main__":
    # Example usage
    service = FileConversionService()
    
    # Print supported formats
    print("Supported formats:")
    for conv_type, formats in service.list_supported_formats().items():
        print(f"{conv_type}: {formats}")

class BatchConverter:
    """Handle batch file conversions with parallel processing"""
    
    def __init__(self, service: FileConversionService):
        self.service = service
    
    def convert_directory(self, input_dir: str, output_dir: str, 
                         input_format: str, output_format: str, 
                         preserve_structure: bool = True) -> Dict[str, bool]:
        """Convert all files of input_format in directory to output_format"""
        input_path = Path(input_dir)
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        results = {}
        
        # Find all files with input format
        pattern = f"**/*.{input_format}" if preserve_structure else f"*.{input_format}"
        for file_path in input_path.glob(pattern):
            if file_path.is_file():
                # Calculate output path
                relative_path = file_path.relative_to(input_path)
                output_file = output_path / relative_path.with_suffix(f'.{output_format}')
                
                # Create output directory if needed
                output_file.parent.mkdir(parents=True, exist_ok=True)
                
                # Convert file
                success = self.service.convert_file(str(file_path), str(output_file))
                results[str(file_path)] = success
                
                if success:
                    print(f"✓ Converted: {file_path} -> {output_file}")
                else:
                    print(f"✗ Failed: {file_path}")
        
        return results

def main():
    """CLI interface for file conversion"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Universal File Converter')
    parser.add_argument('input', nargs='?', help='Input file or directory path')
    parser.add_argument('output', nargs='?', help='Output file or directory path')
    parser.add_argument('--input-format', help='Input format for batch conversion')
    parser.add_argument('--output-format', help='Output format for batch conversion')
    parser.add_argument('--batch', action='store_true', help='Batch convert directory')
    parser.add_argument('--quality', type=int, help='Quality setting for media conversion')
    parser.add_argument('--list-formats', action='store_true', help='List supported formats')
    
    args = parser.parse_args()
    
    service = FileConversionService()
    
    if args.list_formats:
        print("\nSupported formats by converter type:")
        for conv_type, formats in service.list_supported_formats().items():
            print(f"\n{conv_type.upper()}:")
            print(f"  Input:  {', '.join(formats['input'])}")
            print(f"  Output: {', '.join(formats['output'])}")
        return
    
    if args.batch:
        if not args.input_format or not args.output_format:
            print("Error: --input-format and --output-format required for batch conversion")
            return
        
        batch_converter = BatchConverter(service)
        results = batch_converter.convert_directory(
            args.input, args.output, 
            args.input_format, args.output_format
        )
        
        success_count = sum(results.values())
        total_count = len(results)
        print(f"\nBatch conversion complete: {success_count}/{total_count} files converted")
    
    else:
        # Single file conversion
        kwargs = {}
        if args.quality:
            kwargs['quality'] = args.quality
        
        success = service.convert_file(args.input, args.output, **kwargs)
        if success:
            print(f"Successfully converted: {args.input} -> {args.output}")
        else:
            print(f"Conversion failed: {args.input} -> {args.output}")

if __name__ == "__main__":
    main()